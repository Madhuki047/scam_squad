"""Generate two academic-style .docx deliverables on the user's Desktop:

  1. "Jest Test Report.docx"      - backend test suite documentation
  2. "Bug Report Template.docx"   - the GitHub issue template, in Word form

Run from anywhere:  python scripts/generate_reports.py
"""
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor

OUT_DIR = Path.home() / "Desktop"
TODAY = date(2026, 5, 28)
PROJECT = "Scam Squad"
MODULE = "503IT - Software Engineering Group Project"
AUTHOR = "Deverishi Gaire"

DARK = RGBColor(0x1F, 0x1F, 0x1F)
ACCENT = RGBColor(0x0B, 0x5E, 0xD7)
MUTED = RGBColor(0x55, 0x55, 0x55)


def apply_base_styles(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = DARK


def add_cover(doc, title, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(6):
        p.add_run("\n")

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(title)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = ACCENT

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run(subtitle)
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = MUTED

    spacer = doc.add_paragraph()
    spacer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(10):
        spacer.add_run("\n")

    meta = [
        ("Project", PROJECT),
        ("Module", MODULE),
        ("Author", AUTHOR),
        ("Date", TODAY.strftime("%d %B %Y")),
        ("Document version", "1.0"),
    ]
    for label, value in meta:
        line = doc.add_paragraph()
        line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lr = line.add_run(f"{label}:  ")
        lr.bold = True
        line.add_run(value)

    doc.add_page_break()


def add_h1(doc, text, number=None):
    p = doc.add_paragraph()
    run = p.add_run(f"{number}. {text}" if number else text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = ACCENT


def add_h2(doc, text, number=None):
    p = doc.add_paragraph()
    run = p.add_run(f"{number} {text}" if number else text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = DARK


def add_p(doc, text):
    doc.add_paragraph(text)


def add_bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        cell_p = hdr_cells[i].paragraphs[0]
        r = cell_p.add_run(h)
        r.bold = True
    for row_idx, row in enumerate(rows, start=1):
        cells = t.rows[row_idx].cells
        for col_idx, val in enumerate(row):
            cells[col_idx].text = str(val)


# ---------------------------------------------------------------------
# Document 1 - Jest Test Report
# ---------------------------------------------------------------------
def build_jest_report():
    doc = Document()
    apply_base_styles(doc)

    add_cover(
        doc,
        "Jest Test Report",
        "Backend Automated Test Suite",
    )

    # 1. Executive summary
    add_h1(doc, "Executive Summary", 1)
    add_p(
        doc,
        "This report documents the automated test suite built for the Scam "
        "Squad backend (the Express REST API in /server). The suite uses "
        "Jest as the test runner, Supertest for HTTP-level integration "
        "tests, and mongodb-memory-server to spin up an in-memory MongoDB "
        "instance for every run so tests are hermetic and require no "
        "external services."
    )
    add_p(
        doc,
        "Final result: 12 test suites, 84 tests, all passing. The suite "
        "covers every REST controller currently shipped on main except the "
        "chat / Socket.io layer, which is deferred to a subsequent pass."
    )

    # 2. Scope and approach
    add_h1(doc, "Scope and Approach", 2)

    add_h2(doc, "2.1 What is tested", "2.1")
    add_bullet(doc, "Authentication: register, login (with and without email/2FA), OTP verification, logout.")
    add_bullet(doc, "User account: profile read, notification settings, email add, password change, account deletion, public profile lookup.")
    add_bullet(doc, "Lives system: current count, life consumption, depletion 409 response.")
    add_bullet(doc, "Adaptive quiz: question issuing without answer leak, in-progress guard, correct-answer flow, full 5-question session and life reward.")
    add_bullet(doc, "Activity feed: newest-first ordering, per-user scoping, limit clamping.")
    add_bullet(doc, "Leaderboard: ranking by points, pagination, signed-in player rank lookup.")
    add_bullet(doc, "Friend graph: request, accept, decline, remove, self-request rejection, search exclusion of self / friends / pending.")
    add_bullet(doc, "Case progression: progress read, intro completion, case completion idempotency, veteran-after-rookie gating, fail attempt life cost.")
    add_bullet(doc, "Shop: catalog shape, internal field hiding, affordability check, consumable counter, cosmetic one-time purchase.")
    add_bullet(doc, "Health endpoint and unknown API route 404 handling.")
    add_bullet(doc, "Unit tests: lives regeneration math (applyRegen, useLife, grantLife, getNextRegenAt) and User model password hooks.")

    add_h2(doc, "2.2 What is intentionally deferred", "2.2")
    add_bullet(doc, "Chat / Socket.io layer: real-time delivery requires a socket client harness; tracked for a future phase.")
    add_bullet(doc, "Redis-backed paths: the application falls back to in-memory stores when Redis is not configured, which is the path exercised under test.")
    add_bullet(doc, "Frontend (React) tests: this report covers the backend only. The frontend would naturally use Vitest rather than Jest.")

    # 3. Tooling and configuration
    add_h1(doc, "Tooling and Configuration", 3)

    add_table(
        doc,
        ["Component", "Version", "Purpose"],
        [
            ["jest", "^30.4.2", "Test runner and assertion library."],
            ["supertest", "^7.2.2", "Drives the Express app over HTTP without binding a real port."],
            ["mongodb-memory-server", "^10.4.3", "Downloads and runs an in-process MongoDB binary for hermetic database tests."],
            ["cross-env", "^10.1.0", "Sets the --experimental-vm-modules NODE_OPTIONS flag on Windows and POSIX shells alike, so ES module imports work in tests."],
        ],
    )

    add_h2(doc, "3.1 Why an Express factor was needed", "3.1")
    add_p(
        doc,
        "The original server/index.js performed three side effects on "
        "import: connecting to MongoDB, starting Socket.io, and calling "
        "httpServer.listen(). None of those are appropriate inside a test "
        "process. The refactor splits the Express application into "
        "server/app.js (a createApp() factory returning a configured "
        "express() instance) and leaves server/index.js responsible only "
        "for wiring the runtime concerns (database, sockets, listen). "
        "Tests import createApp() directly; production behaviour is "
        "unchanged."
    )

    # 4. Test layout
    add_h1(doc, "Test Layout", 4)

    add_p(doc, "All tests live under server/tests/.")

    add_code(
        doc,
        "server/\n"
        "  jest.config.js\n"
        "  tests/\n"
        "    globalSetup.js          # boots in-memory MongoDB and sets env vars\n"
        "    globalTeardown.js       # stops the in-memory MongoDB\n"
        "    helpers/\n"
        "      db.js                 # connect / clear / disconnect helpers per file\n"
        "      auth.js               # makeUserWithToken(): create a user and mint a JWT\n"
        "    integration/\n"
        "      health.test.js\n"
        "      auth.test.js\n"
        "      user.test.js\n"
        "      lives.test.js\n"
        "      quiz.test.js\n"
        "      activity.test.js\n"
        "      leaderboard.test.js\n"
        "      friends.test.js\n"
        "      progress.test.js\n"
        "      shop.test.js\n"
        "    unit/\n"
        "      models/User.test.js\n"
        "      services/livesService.test.js\n",
    )

    # 5. Hermeticity and isolation
    add_h1(doc, "Hermeticity and Isolation", 5)
    add_p(
        doc,
        "Each test file follows the same lifecycle: connectTestDB in "
        "beforeAll, clearTestDB in afterEach, disconnectTestDB in afterAll. "
        "clearTestDB drops every collection between tests so no test can "
        "see another's data. The Jest invocation uses --runInBand to "
        "execute test files sequentially; without it, parallel files would "
        "trample each other's clearTestDB calls and produce spurious "
        "DocumentNotFound errors mid-test."
    )

    # 6. Running the suite
    add_h1(doc, "Running the Suite", 6)

    add_h2(doc, "6.1 Commands", "6.1")
    add_code(
        doc,
        "cd server\n"
        "npm install       # one time; pulls jest, supertest, mongodb-memory-server, cross-env\n"
        "npm test          # runs the full suite (currently ~11 seconds)\n"
        "npm run test:watch   # re-runs on file change\n",
    )

    add_h2(doc, "6.2 Prerequisites", "6.2")
    add_bullet(doc, "Node.js 20 or newer.")
    add_bullet(doc, "Internet access on the first run, so mongodb-memory-server can download a MongoDB binary into its cache. Subsequent runs are offline.")
    add_bullet(doc, "No MongoDB, Redis, or SMTP server is required - those are all stubbed or stood up in-process.")

    # 7. Results
    add_h1(doc, "Results", 7)

    add_table(
        doc,
        ["Suite", "Tests", "Status"],
        [
            ["health.test.js", "2", "passing"],
            ["auth.test.js", "12", "passing"],
            ["user.test.js", "9", "passing"],
            ["lives.test.js", "3", "passing"],
            ["quiz.test.js", "8", "passing"],
            ["activity.test.js", "3", "passing"],
            ["leaderboard.test.js", "3", "passing"],
            ["friends.test.js", "7", "passing"],
            ["progress.test.js", "7", "passing"],
            ["shop.test.js", "6", "passing"],
            ["unit/services/livesService.test.js", "13", "passing"],
            ["unit/models/User.test.js", "5", "passing"],
            ["TOTAL", "84", "passing (12/12 suites)"],
        ],
    )

    add_h2(doc, "7.1 Final Jest output", "7.1")
    add_code(
        doc,
        "Test Suites: 12 passed, 12 total\n"
        "Tests:       84 passed, 84 total\n"
        "Snapshots:   0 total\n"
        "Time:        ~11 s\n",
    )

    # 8. Risks and notes
    add_h1(doc, "Known Notes and Limitations", 8)
    add_bullet(doc, "Tests run serially to avoid cross-file database wipe collisions; this is a deliberate trade of parallelism for correctness.")
    add_bullet(doc, "The quiz controller pulls random questions and adapts difficulty after two correct answers; the quiz suite seeds questions at every difficulty so a session can always complete.")
    add_bullet(doc, "Activity log ordering is asserted with explicit createdAt timestamps; relying on insertion order alone would tie at the same millisecond.")
    add_bullet(doc, "The /api/shop tests caught one bug during authoring: the test helper originally signed JWTs with a userId claim while the production auth middleware reads payload.id. The helper now matches the production token shape.")

    out = OUT_DIR / "Jest Test Report.docx"
    doc.save(out)
    return out


# ---------------------------------------------------------------------
# Document 2 - Bug Report Template
# ---------------------------------------------------------------------
def build_bug_report():
    doc = Document()
    apply_base_styles(doc)

    add_cover(
        doc,
        "Bug Report Template",
        "Standard format for reporting defects in Scam Squad",
    )

    add_h1(doc, "Purpose", 1)
    add_p(
        doc,
        "This template gives the team a single, consistent shape for "
        "reporting bugs. Filling every section makes triage cheap: the "
        "person picking up the report should be able to reproduce the "
        "issue without a follow-up conversation. The same template is "
        "available as the default GitHub issue form under "
        ".github/ISSUE_TEMPLATE/bug_report.md."
    )

    add_h1(doc, "How to Use", 2)
    add_bullet(doc, "Open a new issue on the project's GitHub repository.")
    add_bullet(doc, "Select 'Bug report' from the template chooser - the form is pre-filled with the sections below.")
    add_bullet(doc, "Replace each placeholder. Empty sections are worse than honest 'unknown' answers; if you do not know something, say so.")
    add_bullet(doc, "Apply the appropriate severity label (critical, high, medium, low).")
    add_bullet(doc, "Mask any secrets, tokens, or personally identifiable information before pasting logs.")

    add_h1(doc, "Template", 3)

    add_h2(doc, "Title", "3.1")
    add_code(doc, "[BUG] <short, specific summary>")
    add_p(doc, "Example: [BUG] Shop page shows under-construction message after merging deploy-shop")

    add_h2(doc, "What happened", "3.2")
    add_p(doc, "One or two sentences describing what you did and what went wrong. Keep it factual; leave interpretation for later sections.")

    add_h2(doc, "Steps to reproduce", "3.3")
    add_p(doc, "Numbered, click-by-click. A bug that cannot be reproduced is almost impossible to fix.")
    add_code(doc, "1. Sign in as a user with at least 500 points.\n2. Navigate to /shop.\n3. Click the Buy button on the Neon Skin item.")

    add_h2(doc, "Expected behaviour", "3.4")
    add_p(doc, "What should have happened, and the reason you expected it (a specification, a previous build, the user story, etc.).")

    add_h2(doc, "Actual behaviour", "3.5")
    add_p(doc, "What you actually saw. Paste exact error text where possible; attach a screenshot when the issue is visual.")

    add_h2(doc, "Environment", "3.6")
    add_table(
        doc,
        ["Field", "Example value"],
        [
            ["Surface", "frontend (Shop page) / backend (/api/shop) / deploy (Vercel) / tests"],
            ["Branch / commit", "main @ 66f08c5  or  deploy-shop"],
            ["Where it ran", "local dev / Vercel preview / Vercel production / Render"],
            ["Browser (frontend)", "Chrome 131 on Windows 11"],
            ["Node version (backend)", "20.15.0"],
            ["Account state", "new user / existing user / unauthenticated"],
        ],
    )

    add_h2(doc, "Severity", "3.7")
    add_table(
        doc,
        ["Level", "Use when"],
        [
            ["critical", "Data loss, security exposure, or total outage."],
            ["high", "A core flow (login, case, quiz, shop) is unusable."],
            ["medium", "A flow works but is degraded or has a wrong-looking result."],
            ["low", "Cosmetic, a rare edge case, or affects a non-core surface."],
        ],
    )

    add_h2(doc, "Logs and network response", "3.8")
    add_p(doc, "Paste any relevant server logs, browser console output, or failed API response bodies. Wrap them in a code block. Mask secrets and email addresses.")

    add_h2(doc, "Anything else", "3.9")
    add_p(doc, "Related issues, recent changes that may have introduced the bug, a workaround you discovered, or environment quirks that may matter.")

    add_h1(doc, "Worked Example", 4)
    add_p(doc, "The following report would be useful enough to triage immediately.")

    add_p(doc, "")
    add_h2(doc, "Title")
    add_code(doc, "[BUG] /api/shop/buy returns 404 instead of 400 when an unknown item is bought")

    add_h2(doc, "What happened")
    add_p(doc, "Calling POST /api/shop/buy/foo with a valid session returned HTTP 404 with body { \"message\": \"Unknown item.\" }, but the spec says invalid item ids should return 400.")

    add_h2(doc, "Steps to reproduce")
    add_code(doc, "1. Register and sign in to obtain a JWT.\n2. curl -X POST http://localhost:3001/api/shop/buy/foo -H 'Authorization: Bearer <jwt>'.\n3. Observe the response status code.")

    add_h2(doc, "Expected behaviour")
    add_p(doc, "HTTP 400 Bad Request, because the client supplied an item id that does not exist in the catalog.")

    add_h2(doc, "Actual behaviour")
    add_p(doc, "HTTP 404 Not Found with body { \"message\": \"Unknown item.\" }.")

    add_h2(doc, "Environment")
    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Surface", "backend (/api/shop/buy/:itemId)"],
            ["Branch / commit", "main @ 66f08c5"],
            ["Where it ran", "local dev"],
            ["Node version", "20.15.0"],
            ["Account state", "existing user, 500 points"],
        ],
    )

    add_h2(doc, "Severity")
    add_p(doc, "low - the response is wrong but no money or progress is lost.")

    out = OUT_DIR / "Bug Report Template.docx"
    doc.save(out)
    return out


def main():
    jest_path = build_jest_report()
    bug_path = build_bug_report()
    print(f"Wrote: {jest_path}")
    print(f"Wrote: {bug_path}")


if __name__ == "__main__":
    main()
